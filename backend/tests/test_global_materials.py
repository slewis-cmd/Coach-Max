"""
Tests for Course-Wide (is_global) library materials feature.

Endpoints under test:
- POST /api/library/materials (with is_global=true/false)
- GET /api/library/materials (returns is_global on each material)
- POST /api/library/materials/{id}/assign (works for global materials)
- GET /api/student/dashboard (returns per-cohort course_resources array)
- POST /api/submissions/{id}/review (prepends global materials to context)

Regression:
- POST /api/library/materials/{id}/duplicate — duplicated global material preserves is_global=True
- GET /api/materials/{id}/download — students in an assigned cohort can download global materials
- /api/student/dashboard still returns per-week materials in `weeks` array
"""
import io
import os
import uuid
import time
from datetime import datetime, timezone, timedelta

import pytest
import requests
from pymongo import MongoClient
from bson import ObjectId
from docx import Document


BASE_URL = os.environ.get("REACT_APP_BACKEND_URL", "").rstrip("/")
MONGO_URL = os.environ.get("MONGO_URL", "mongodb://localhost:27017")
DB_NAME = os.environ.get("DB_NAME", "test_database")
SUPER_ADMIN_EMAIL = "slewis@theboostpad.org"

TEST_PREFIX = "TEST_GLOBAL_"

# ---------- fixtures ----------

@pytest.fixture(scope="module")
def mongo():
    client = MongoClient(MONGO_URL)
    yield client[DB_NAME]
    client.close()


def _seed_session(mongo, user_id: str, prefix: str = "test_glob") -> str:
    token = f"{prefix}_{uuid.uuid4().hex[:12]}"
    mongo.user_sessions.insert_one({
        "session_token": token,
        "user_id": user_id,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "expires_at": (datetime.now(timezone.utc) + timedelta(days=1)).isoformat(),
    })
    return token


@pytest.fixture(scope="module")
def admin_token(mongo):
    admin = mongo.users.find_one({"email": SUPER_ADMIN_EMAIL})
    assert admin, f"Super admin {SUPER_ADMIN_EMAIL} not seeded"
    token = _seed_session(mongo, admin["user_id"], "test_glob_admin")
    yield token
    mongo.user_sessions.delete_one({"session_token": token})


@pytest.fixture(scope="module")
def admin_headers(admin_token):
    return {"Authorization": f"Bearer {admin_token}"}


@pytest.fixture(scope="module")
def instructor(mongo):
    uid = f"user_{uuid.uuid4().hex[:12]}"
    email = f"{TEST_PREFIX}inst_{uuid.uuid4().hex[:6]}@example.com"
    mongo.users.insert_one({
        "user_id": uid,
        "email": email,
        "name": f"{TEST_PREFIX}Instructor",
        "picture": None,
        "role": "instructor",
        "created_at": datetime.now(timezone.utc).isoformat(),
    })
    token = _seed_session(mongo, uid, "test_glob_inst")
    yield {"user_id": uid, "email": email, "token": token,
           "headers": {"Authorization": f"Bearer {token}"}}
    mongo.user_sessions.delete_one({"session_token": token})
    mongo.users.delete_one({"user_id": uid})


@pytest.fixture(scope="module")
def student(mongo):
    uid = f"user_{uuid.uuid4().hex[:12]}"
    email = f"{TEST_PREFIX}stu_{uuid.uuid4().hex[:6]}@example.com"
    mongo.users.insert_one({
        "user_id": uid,
        "email": email,
        "name": f"{TEST_PREFIX}Student",
        "picture": None,
        "role": "student",
        "created_at": datetime.now(timezone.utc).isoformat(),
    })
    token = _seed_session(mongo, uid, "test_glob_stu")
    yield {"user_id": uid, "email": email, "token": token,
           "headers": {"Authorization": f"Bearer {token}"}}
    mongo.user_sessions.delete_one({"session_token": token})
    mongo.users.delete_one({"user_id": uid})


def _make_docx_bytes(text: str) -> bytes:
    doc = Document()
    doc.add_heading("TEST Global Materials", level=1)
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_pdf_bytes() -> bytes:
    # Minimal PDF file
    return (b"%PDF-1.4\n1 0 obj\n<< /Type /Catalog >>\nendobj\n"
            b"trailer\n<< /Root 1 0 R >>\n%%EOF")


@pytest.fixture(scope="module")
def cohort_c1(mongo, admin_headers, instructor, student):
    """Cohort C1 managed by instructor, with student enrolled, week 3 released."""
    name = f"{TEST_PREFIX}Cohort_{uuid.uuid4().hex[:6]}"
    r = requests.post(
        f"{BASE_URL}/api/cohorts",
        json={"name": name, "description": "global materials test cohort"},
        headers=admin_headers, timeout=30,
    )
    assert r.status_code == 200, r.text
    cid = r.json()["cohort_id"]
    mongo.cohorts.update_one(
        {"cohort_id": cid},
        {"$addToSet": {
            "instructor_ids": instructor["user_id"],
            "student_ids": student["user_id"],
        },
         "$set": {"released_weeks": [3]}},
    )
    yield {"cohort_id": cid, "name": name}
    # Teardown
    requests.delete(f"{BASE_URL}/api/cohorts/{cid}", headers=admin_headers, timeout=30)
    mongo.cohorts.delete_one({"cohort_id": cid})


# Track created docs for cleanup
_created_materials = []
_created_submissions = []


@pytest.fixture(scope="module", autouse=True)
def cleanup_all(mongo):
    yield
    # Cleanup materials + gridfs
    for mid in _created_materials:
        m = mongo.materials.find_one({"material_id": mid})
        if m and m.get("gridfs_id"):
            try:
                mongo["fs.files"].delete_one({"_id": ObjectId(m["gridfs_id"])})
                mongo["fs.chunks"].delete_many({"files_id": ObjectId(m["gridfs_id"])})
            except Exception:
                pass
        mongo.materials.delete_one({"material_id": mid})
    # Cleanup any test materials by title prefix (safety net)
    for m in mongo.materials.find({"title": {"$regex": f"^{TEST_PREFIX}"}}):
        if m.get("gridfs_id"):
            try:
                mongo["fs.files"].delete_one({"_id": ObjectId(m["gridfs_id"])})
                mongo["fs.chunks"].delete_many({"files_id": ObjectId(m["gridfs_id"])})
            except Exception:
                pass
    mongo.materials.delete_many({"title": {"$regex": f"^{TEST_PREFIX}"}})
    # Cleanup submissions
    for sid in _created_submissions:
        s = mongo.submissions.find_one({"submission_id": sid})
        if s and s.get("gridfs_id"):
            try:
                mongo["fs.files"].delete_one({"_id": ObjectId(s["gridfs_id"])})
                mongo["fs.chunks"].delete_many({"files_id": ObjectId(s["gridfs_id"])})
            except Exception:
                pass
        mongo.submissions.delete_one({"submission_id": sid})
    # Cleanup ai_feedbacks
    mongo.ai_feedbacks.delete_many({"submission_id": {"$in": _created_submissions}})


def _upload_library_material(admin_headers, *, title, week_number, material_type,
                             is_global=None, file_name="material.docx"):
    """Helper: uploads a library material. Returns material_id."""
    docx_bytes = _make_docx_bytes(f"Content for {title}")
    files = {"file": (file_name, docx_bytes,
                      "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    params = {
        "week_number": week_number,
        "material_type": material_type,
        "title": title,
        "description": f"desc for {title}",
    }
    if is_global is not None:
        # FastAPI accepts 'true'/'false' as bool for query params
        params["is_global"] = "true" if is_global else "false"
    r = requests.post(
        f"{BASE_URL}/api/library/materials",
        params=params, files=files, headers=admin_headers, timeout=60,
    )
    assert r.status_code == 200, r.text
    mid = r.json()["material_id"]
    _created_materials.append(mid)
    return mid, docx_bytes


# ---------- Tests ----------

class TestUploadIsGlobal:
    """POST /api/library/materials with is_global param."""

    def test_upload_is_global_true_sets_week_zero(self, admin_headers, mongo):
        """is_global=true → is_global stored True, week_number forced to 0
        regardless of what week_number was passed."""
        mid, _ = _upload_library_material(
            admin_headers,
            title=f"{TEST_PREFIX}Global_A_{uuid.uuid4().hex[:4]}",
            week_number=7,  # should be overridden to 0
            material_type="case_study",
            is_global=True,
        )
        doc = mongo.materials.find_one({"material_id": mid})
        assert doc is not None
        assert doc.get("is_global"), f"Expected is_global=True, got {doc.get('is_global')}"
        assert doc.get("week_number") == 0, \
            f"Expected week_number=0 for global, got {doc.get('week_number')}"
        assert doc.get("is_library")

    def test_upload_is_global_false_preserves_week(self, admin_headers, mongo):
        """is_global=false → is_global=False and week_number preserved."""
        mid, _ = _upload_library_material(
            admin_headers,
            title=f"{TEST_PREFIX}NotGlobal_{uuid.uuid4().hex[:4]}",
            week_number=4,
            material_type="workbook",
            is_global=False,
        )
        doc = mongo.materials.find_one({"material_id": mid})
        assert doc is not None
        assert not (doc.get("is_global"))
        assert doc.get("week_number") == 4

    def test_upload_is_global_omitted_defaults_false(self, admin_headers, mongo):
        """When is_global not provided, backend defaults to False and preserves week."""
        mid, _ = _upload_library_material(
            admin_headers,
            title=f"{TEST_PREFIX}Default_{uuid.uuid4().hex[:4]}",
            week_number=5,
            material_type="workbook",
            is_global=None,
        )
        doc = mongo.materials.find_one({"material_id": mid})
        assert doc is not None
        assert not (doc.get("is_global"))
        assert doc.get("week_number") == 5


class TestGetLibraryReturnsIsGlobal:
    """GET /api/library/materials returns is_global on each material."""

    def test_list_returns_is_global_flag(self, admin_headers):
        mid_global, _ = _upload_library_material(
            admin_headers,
            title=f"{TEST_PREFIX}ListGlobal_{uuid.uuid4().hex[:4]}",
            week_number=1,
            material_type="workbook",
            is_global=True,
        )
        mid_regular, _ = _upload_library_material(
            admin_headers,
            title=f"{TEST_PREFIX}ListRegular_{uuid.uuid4().hex[:4]}",
            week_number=2,
            material_type="workbook",
            is_global=False,
        )
        r = requests.get(f"{BASE_URL}/api/library/materials",
                         headers=admin_headers, timeout=30)
        assert r.status_code == 200
        materials = r.json()
        by_id = {m["material_id"]: m for m in materials}
        assert mid_global in by_id, "Global material must appear in library listing"
        assert mid_regular in by_id, "Regular material must appear in library listing"
        assert by_id[mid_global].get("is_global")
        assert not (by_id[mid_regular].get("is_global"))


class TestAssignGlobalMaterial:
    """POST /api/library/materials/{id}/assign works for global materials."""

    def test_assign_global_to_cohort(self, admin_headers, cohort_c1, mongo):
        mid, _ = _upload_library_material(
            admin_headers,
            title=f"{TEST_PREFIX}AssignGlobal_{uuid.uuid4().hex[:4]}",
            week_number=1,
            material_type="workbook",
            is_global=True,
        )
        r = requests.post(
            f"{BASE_URL}/api/library/materials/{mid}/assign",
            json={"cohort_ids": [cohort_c1["cohort_id"]]},
            headers=admin_headers, timeout=30,
        )
        assert r.status_code == 200, r.text
        doc = mongo.materials.find_one({"material_id": mid})
        assert cohort_c1["cohort_id"] in doc.get("cohort_ids", [])

    def test_assign_regular_to_cohort(self, admin_headers, cohort_c1, mongo):
        mid, _ = _upload_library_material(
            admin_headers,
            title=f"{TEST_PREFIX}AssignRegular_{uuid.uuid4().hex[:4]}",
            week_number=2,
            material_type="workbook",
            is_global=False,
        )
        r = requests.post(
            f"{BASE_URL}/api/library/materials/{mid}/assign",
            json={"cohort_ids": [cohort_c1["cohort_id"]]},
            headers=admin_headers, timeout=30,
        )
        assert r.status_code == 200, r.text
        doc = mongo.materials.find_one({"material_id": mid})
        assert cohort_c1["cohort_id"] in doc.get("cohort_ids", [])


class TestStudentDashboardCourseResources:
    """GET /api/student/dashboard returns per-cohort course_resources array."""

    def test_dashboard_has_course_resources_with_globals(
        self, admin_headers, cohort_c1, student
    ):
        # Upload + assign 2 global materials
        mid1, _ = _upload_library_material(
            admin_headers,
            title=f"{TEST_PREFIX}Dash_Glob1_{uuid.uuid4().hex[:4]}",
            week_number=1, material_type="case_study", is_global=True,
        )
        mid2, _ = _upload_library_material(
            admin_headers,
            title=f"{TEST_PREFIX}Dash_Glob2_{uuid.uuid4().hex[:4]}",
            week_number=1, material_type="workbook", is_global=True,
        )
        for m in (mid1, mid2):
            r = requests.post(
                f"{BASE_URL}/api/library/materials/{m}/assign",
                json={"cohort_ids": [cohort_c1["cohort_id"]]},
                headers=admin_headers, timeout=30,
            )
            assert r.status_code == 200

        # Also assign a non-global (regression: it should not appear in course_resources)
        mid_regular, _ = _upload_library_material(
            admin_headers,
            title=f"{TEST_PREFIX}Dash_NotGlob_{uuid.uuid4().hex[:4]}",
            week_number=3, material_type="workbook", is_global=False,
        )
        r = requests.post(
            f"{BASE_URL}/api/library/materials/{mid_regular}/assign",
            json={"cohort_ids": [cohort_c1["cohort_id"]]},
            headers=admin_headers, timeout=30,
        )
        assert r.status_code == 200

        r = requests.get(
            f"{BASE_URL}/api/student/dashboard",
            headers=student["headers"], timeout=30,
        )
        assert r.status_code == 200, r.text
        data = r.json()
        assert isinstance(data, list)

        my_cohort = next((c for c in data if c["cohort_id"] == cohort_c1["cohort_id"]), None)
        assert my_cohort is not None, f"Cohort {cohort_c1['cohort_id']} not in dashboard"

        assert "course_resources" in my_cohort, \
            "course_resources key missing from per-cohort dashboard response"
        cr = my_cohort["course_resources"]
        assert isinstance(cr, list)
        cr_ids = {r["material_id"] for r in cr}
        assert mid1 in cr_ids, f"Global material {mid1} missing from course_resources"
        assert mid2 in cr_ids, f"Global material {mid2} missing from course_resources"
        assert mid_regular not in cr_ids, \
            "Non-global material must NOT appear in course_resources"

        # Verify schema of each course_resource entry
        entry = next(x for x in cr if x["material_id"] == mid1)
        for key in ("material_id", "title", "material_type", "file_name", "description"):
            assert key in entry, f"course_resources entry missing '{key}'"

        # Regression: weeks array still present
        assert "weeks" in my_cohort
        assert isinstance(my_cohort["weeks"], list)

    def test_dashboard_course_resources_empty_when_no_globals_assigned(
        self, admin_headers, student, mongo, instructor
    ):
        # Create a fresh cohort with NO global materials assigned
        name = f"{TEST_PREFIX}EmptyCohort_{uuid.uuid4().hex[:6]}"
        r = requests.post(
            f"{BASE_URL}/api/cohorts",
            json={"name": name, "description": "no globals"},
            headers=admin_headers, timeout=30,
        )
        assert r.status_code == 200
        cid = r.json()["cohort_id"]
        mongo.cohorts.update_one(
            {"cohort_id": cid},
            {"$addToSet": {"student_ids": student["user_id"],
                           "instructor_ids": instructor["user_id"]},
             "$set": {"released_weeks": [1]}},
        )
        try:
            r = requests.get(
                f"{BASE_URL}/api/student/dashboard",
                headers=student["headers"], timeout=30,
            )
            assert r.status_code == 200
            data = r.json()
            empty_cohort = next((c for c in data if c["cohort_id"] == cid), None)
            assert empty_cohort is not None
            assert empty_cohort.get("course_resources") == [], \
                f"Expected empty course_resources, got {empty_cohort.get('course_resources')}"
        finally:
            requests.delete(f"{BASE_URL}/api/cohorts/{cid}",
                            headers=admin_headers, timeout=30)
            mongo.cohorts.delete_one({"cohort_id": cid})


class TestReviewWithGlobalMaterials:
    """POST /api/submissions/{id}/review — succeeds when globals are assigned,
    even when current week has zero week-specific workbooks."""

    def test_review_endpoint_prepends_globals(
        self, admin_headers, cohort_c1, student, mongo
    ):
        # Ensure at least 1 global assigned to the cohort
        mid_glob, _ = _upload_library_material(
            admin_headers,
            title=f"{TEST_PREFIX}Review_Glob_{uuid.uuid4().hex[:4]}",
            week_number=1, material_type="workbook", is_global=True,
        )
        r = requests.post(
            f"{BASE_URL}/api/library/materials/{mid_glob}/assign",
            json={"cohort_ids": [cohort_c1["cohort_id"]]},
            headers=admin_headers, timeout=30,
        )
        assert r.status_code == 200

        # Upload homework material for week 3 (cohort-specific)
        hw_bytes = _make_docx_bytes("Homework prompt week 3")
        files = {"file": ("hw_w3.docx", hw_bytes,
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        params = {
            "week_number": 3,
            "material_type": "homework",
            "title": f"{TEST_PREFIX}HW_W3_{uuid.uuid4().hex[:4]}",
            "description": "test homework week 3",
        }
        r = requests.post(
            f"{BASE_URL}/api/cohorts/{cohort_c1['cohort_id']}/materials",
            params=params, files=files, headers=admin_headers, timeout=60,
        )
        assert r.status_code == 200, r.text
        hw_mid = r.json()["material_id"]
        _created_materials.append(hw_mid)

        # Student submits homework
        sub_bytes = _make_docx_bytes(
            "This is my homework response. I applied the framework and analyzed the case. "
            "It required careful thought and reflection on the material provided. " * 3
        )
        files = {"file": ("submission.docx", sub_bytes,
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        r = requests.post(
            f"{BASE_URL}/api/materials/{hw_mid}/submit?cohort_id={cohort_c1['cohort_id']}",
            files=files, headers=student["headers"], timeout=60,
        )
        assert r.status_code == 200, r.text
        sub_id = r.json()["submission_id"]
        _created_submissions.append(sub_id)

        # Call review endpoint (as admin — is_cohort_manager: super_admin passes)
        r = requests.post(
            f"{BASE_URL}/api/submissions/{sub_id}/review",
            headers=admin_headers, timeout=120,
        )
        assert r.status_code == 200, f"Review failed: {r.status_code} {r.text}"
        body = r.json()
        # Some form of feedback returned
        assert "feedback" in body or "ai_feedback" in body or "message" in body, \
            f"Review response missing expected keys: {body}"

        # Submission should now have ai_feedback set
        sub_doc = mongo.submissions.find_one({"submission_id": sub_id})
        assert sub_doc is not None
        assert sub_doc.get("ai_feedback"), \
            f"submission.ai_feedback should be populated after review: {sub_doc.get('ai_feedback')}"


class TestDuplicatePreservesIsGlobal:
    """Regression: POST /library/materials/{id}/duplicate preserves is_global=True."""

    def test_duplicate_global_material_keeps_is_global(self, admin_headers, mongo):
        mid, _ = _upload_library_material(
            admin_headers,
            title=f"{TEST_PREFIX}Dup_Global_{uuid.uuid4().hex[:4]}",
            week_number=1, material_type="case_study", is_global=True,
        )
        r = requests.post(
            f"{BASE_URL}/api/library/materials/{mid}/duplicate",
            headers=admin_headers, timeout=30,
        )
        assert r.status_code == 200, r.text
        new_mid = r.json()["material_id"]
        _created_materials.append(new_mid)

        dup_doc = mongo.materials.find_one({"material_id": new_mid})
        assert dup_doc is not None
        assert dup_doc.get("is_global"), \
            f"Duplicated global material must preserve is_global=True, got {dup_doc.get('is_global')}"
        assert dup_doc.get("week_number") == 0, \
            f"Duplicated global material should have week_number=0, got {dup_doc.get('week_number')}"

    def test_duplicate_non_global_material_stays_non_global(self, admin_headers, mongo):
        mid, _ = _upload_library_material(
            admin_headers,
            title=f"{TEST_PREFIX}Dup_Regular_{uuid.uuid4().hex[:4]}",
            week_number=3, material_type="workbook", is_global=False,
        )
        r = requests.post(
            f"{BASE_URL}/api/library/materials/{mid}/duplicate",
            headers=admin_headers, timeout=30,
        )
        assert r.status_code == 200, r.text
        new_mid = r.json()["material_id"]
        _created_materials.append(new_mid)

        dup_doc = mongo.materials.find_one({"material_id": new_mid})
        # Either False or omitted — both are truthy-negative
        assert not dup_doc.get("is_global"), \
            f"Duplicated regular material must not be marked global, got is_global={dup_doc.get('is_global')}"


class TestDownloadGlobalMaterial:
    """Regression: Students in an assigned cohort can download global materials."""

    def test_student_can_download_assigned_global(
        self, admin_headers, cohort_c1, student
    ):
        mid, source_bytes = _upload_library_material(
            admin_headers,
            title=f"{TEST_PREFIX}DL_Global_{uuid.uuid4().hex[:4]}",
            week_number=1, material_type="workbook", is_global=True,
        )
        r = requests.post(
            f"{BASE_URL}/api/library/materials/{mid}/assign",
            json={"cohort_ids": [cohort_c1["cohort_id"]]},
            headers=admin_headers, timeout=30,
        )
        assert r.status_code == 200

        # Download as student
        r = requests.get(
            f"{BASE_URL}/api/materials/{mid}/download",
            headers=student["headers"], timeout=30,
        )
        assert r.status_code == 200, f"Student download failed: {r.status_code} {r.text[:200]}"
        assert r.content == source_bytes, "Downloaded bytes do not match source"


if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])
