"""
Tests for inline preview of student submissions.

Covers:
  - GET /api/submissions/{id}/download?inline=1 -> Content-Disposition: inline + application/pdf for PDFs
  - GET /api/submissions/{id}/download          -> Content-Disposition: attachment (back-compat)
  - GET /api/submissions/{id}/preview-text      -> returns {text, file_name} with extracted text
  - preview-text access control (student -> 403 on other student's submission; wrong instructor -> 403)
  - preview-text 410 for legacy file_path-only submissions
  - Regression: GET /api/submissions/{id} still returns submission with material + student
"""
import io
import os
import uuid
from datetime import datetime, timezone, timedelta

import pytest
import requests
from pymongo import MongoClient
from docx import Document
from fpdf import FPDF


BASE_URL = os.environ.get("REACT_APP_BACKEND_URL", "").rstrip("/")
MONGO_URL = os.environ.get("MONGO_URL", "mongodb://localhost:27017")
DB_NAME = os.environ.get("DB_NAME", "test_database")
SUPER_ADMIN_EMAIL = "slewis@theboostpad.org"


# ---------- helpers ----------

def _make_docx_bytes(text: str) -> bytes:
    doc = Document()
    doc.add_heading("TEST Preview Homework", level=1)
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_pdf_bytes(text: str) -> bytes:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.cell(0, 10, text)
    raw = pdf.output(dest="S")
    if isinstance(raw, str):
        return raw.encode("latin-1")
    if isinstance(raw, bytearray):
        return bytes(raw)
    return raw


# ---------- fixtures ----------

@pytest.fixture(scope="module")
def mongo():
    client = MongoClient(MONGO_URL)
    yield client[DB_NAME]
    client.close()


def _mk_session(mongo, user_id: str, prefix: str = "test_prev") -> str:
    token = f"{prefix}_{uuid.uuid4().hex[:12]}"
    mongo.user_sessions.insert_one({
        "session_token": token,
        "user_id": user_id,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "expires_at": (datetime.now(timezone.utc) + timedelta(days=1)).isoformat(),
    })
    return token


def _mk_student(mongo) -> dict:
    uid = f"user_{uuid.uuid4().hex[:12]}"
    email = f"TEST_prev_{uuid.uuid4().hex[:6]}@example.com"
    mongo.users.insert_one({
        "user_id": uid,
        "email": email,
        "name": "TEST Preview Student",
        "picture": None,
        "role": "student",
        "created_at": datetime.now(timezone.utc).isoformat(),
    })
    token = _mk_session(mongo, uid, "test_prev_stu")
    return {"user_id": uid, "email": email, "token": token}


def _mk_instructor(mongo) -> dict:
    uid = f"user_{uuid.uuid4().hex[:12]}"
    email = f"TEST_prev_inst_{uuid.uuid4().hex[:6]}@example.com"
    mongo.users.insert_one({
        "user_id": uid,
        "email": email,
        "name": "TEST Preview Instructor",
        "picture": None,
        "role": "instructor",
        "created_at": datetime.now(timezone.utc).isoformat(),
    })
    token = _mk_session(mongo, uid, "test_prev_inst")
    return {"user_id": uid, "email": email, "token": token}


@pytest.fixture(scope="module")
def admin_token(mongo):
    admin = mongo.users.find_one({"email": SUPER_ADMIN_EMAIL})
    assert admin, f"Super admin {SUPER_ADMIN_EMAIL} not seeded"
    token = _mk_session(mongo, admin["user_id"], "test_prev_admin")
    yield token
    mongo.user_sessions.delete_one({"session_token": token})


@pytest.fixture(scope="module")
def auth_headers(admin_token):
    return {"Authorization": f"Bearer {admin_token}"}


@pytest.fixture(scope="module")
def student_a(mongo):
    s = _mk_student(mongo)
    yield s
    mongo.user_sessions.delete_one({"session_token": s["token"]})
    mongo.users.delete_one({"user_id": s["user_id"]})


@pytest.fixture(scope="module")
def student_b(mongo):
    s = _mk_student(mongo)
    yield s
    mongo.user_sessions.delete_one({"session_token": s["token"]})
    mongo.users.delete_one({"user_id": s["user_id"]})


@pytest.fixture(scope="module")
def outsider_instructor(mongo):
    """An instructor NOT assigned to the test cohort."""
    i = _mk_instructor(mongo)
    yield i
    mongo.user_sessions.delete_one({"session_token": i["token"]})
    mongo.users.delete_one({"user_id": i["user_id"]})


@pytest.fixture(scope="module")
def test_cohort(mongo, auth_headers, student_a, student_b):
    name = f"TEST_Preview_Cohort_{uuid.uuid4().hex[:6]}"
    r = requests.post(
        f"{BASE_URL}/api/cohorts",
        json={"name": name, "description": "preview endpoint test"},
        headers=auth_headers, timeout=30,
    )
    assert r.status_code == 200, r.text
    cohort_id = r.json()["cohort_id"]
    # Enroll both students
    mongo.cohorts.update_one(
        {"cohort_id": cohort_id},
        {"$addToSet": {"student_ids": {"$each": [student_a["user_id"], student_b["user_id"]]}}},
    )
    yield cohort_id
    requests.delete(f"{BASE_URL}/api/cohorts/{cohort_id}", headers=auth_headers, timeout=30)


@pytest.fixture(scope="module")
def library_material(mongo, auth_headers, test_cohort):
    docx_bytes = _make_docx_bytes("Preview assignment reference")
    files = {"file": ("assignment.docx", docx_bytes,
                      "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    params = {
        "week_number": 1,
        "material_type": "homework",
        "title": f"TEST Preview HW {uuid.uuid4().hex[:4]}",
        "description": "automated preview test material",
    }
    r = requests.post(f"{BASE_URL}/api/library/materials", params=params, files=files,
                      headers=auth_headers, timeout=60)
    assert r.status_code == 200, r.text
    material_id = r.json()["material_id"]
    r2 = requests.post(f"{BASE_URL}/api/library/materials/{material_id}/assign",
                       json={"cohort_ids": [test_cohort]}, headers=auth_headers, timeout=30)
    assert r2.status_code == 200, r2.text
    r3 = requests.post(f"{BASE_URL}/api/cohorts/{test_cohort}/release-week",
                       json={"week_number": 1}, headers=auth_headers, timeout=30)
    assert r3.status_code == 200, r3.text
    yield material_id
    requests.delete(f"{BASE_URL}/api/library/materials/{material_id}",
                    headers=auth_headers, timeout=30)


def _submit(student, material_id, cohort_id, filename, content_bytes, mime):
    files = {"file": (filename, content_bytes, mime)}
    r = requests.post(
        f"{BASE_URL}/api/materials/{material_id}/submit",
        files=files, params={"cohort_id": cohort_id},
        headers={"Authorization": f"Bearer {student['token']}"}, timeout=60,
    )
    assert r.status_code == 200, r.text
    return r.json()


@pytest.fixture(scope="module")
def docx_submission(student_a, library_material, test_cohort, mongo):
    body = "This is the student's DOCX submission. " + ("leadership reflection. " * 5)
    _submit(student_a, library_material, test_cohort,
            "my_homework.docx", _make_docx_bytes(body),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    doc = mongo.submissions.find_one({
        "material_id": library_material, "student_id": student_a["user_id"]
    })
    assert doc
    return {"submission_id": doc["submission_id"], "body_text": body}


@pytest.fixture(scope="module")
def pdf_submission(student_b, library_material, test_cohort, mongo):
    pdf_bytes = _make_pdf_bytes("Hello PDF preview content")
    _submit(student_b, library_material, test_cohort,
            "my_homework.pdf", pdf_bytes, "application/pdf")
    doc = mongo.submissions.find_one({
        "material_id": library_material, "student_id": student_b["user_id"]
    })
    assert doc
    return {"submission_id": doc["submission_id"], "bytes": pdf_bytes}


# ---------- tests ----------


class TestInlineDownloadHeaders:
    """?inline=1 query param must set Content-Disposition: inline; PDFs must use application/pdf."""

    def test_pdf_download_default_is_attachment(self, pdf_submission, auth_headers):
        r = requests.get(
            f"{BASE_URL}/api/submissions/{pdf_submission['submission_id']}/download",
            headers=auth_headers, timeout=30,
        )
        assert r.status_code == 200, r.text
        cd = r.headers.get("Content-Disposition", "")
        assert cd.lower().startswith("attachment"), f"expected attachment, got: {cd!r}"
        assert r.content == pdf_submission["bytes"], "bytes should round-trip"

    def test_pdf_download_inline_returns_inline_and_pdf_media_type(self, pdf_submission, auth_headers):
        r = requests.get(
            f"{BASE_URL}/api/submissions/{pdf_submission['submission_id']}/download",
            params={"inline": 1}, headers=auth_headers, timeout=30,
        )
        assert r.status_code == 200, r.text
        cd = r.headers.get("Content-Disposition", "")
        assert cd.lower().startswith("inline"), f"expected inline disposition, got: {cd!r}"
        ct = r.headers.get("Content-Type", "")
        assert ct.startswith("application/pdf"), f"expected application/pdf, got: {ct!r}"
        assert r.content == pdf_submission["bytes"]

    def test_docx_inline_disposition_but_octet_media_type(self, docx_submission, auth_headers):
        # Non-PDFs: Content-Disposition still inline when inline=1, media type is octet-stream.
        r = requests.get(
            f"{BASE_URL}/api/submissions/{docx_submission['submission_id']}/download",
            params={"inline": 1}, headers=auth_headers, timeout=30,
        )
        assert r.status_code == 200, r.text
        cd = r.headers.get("Content-Disposition", "")
        assert cd.lower().startswith("inline"), f"expected inline, got: {cd!r}"
        ct = r.headers.get("Content-Type", "")
        assert ct.startswith("application/octet-stream"), f"expected octet-stream for docx, got: {ct!r}"


class TestPreviewTextEndpoint:
    """GET /preview-text returns extracted text + file_name."""

    def test_docx_preview_text_contains_body(self, docx_submission, auth_headers):
        r = requests.get(
            f"{BASE_URL}/api/submissions/{docx_submission['submission_id']}/preview-text",
            headers=auth_headers, timeout=30,
        )
        assert r.status_code == 200, r.text
        data = r.json()
        assert "text" in data and "file_name" in data
        assert data["file_name"].endswith(".docx")
        assert "leadership reflection" in data["text"].lower(), \
            f"extracted text should include body content; got: {data['text'][:200]!r}"

    def test_pdf_preview_text_returns_text_field(self, pdf_submission, auth_headers):
        r = requests.get(
            f"{BASE_URL}/api/submissions/{pdf_submission['submission_id']}/preview-text",
            headers=auth_headers, timeout=30,
        )
        assert r.status_code == 200, r.text
        data = r.json()
        assert "text" in data and "file_name" in data
        assert isinstance(data["text"], str)
        assert data["file_name"].endswith(".pdf")


class TestPreviewTextAccessControl:
    """preview-text must enforce the same ACL as download."""

    def test_student_cannot_preview_other_students_submission(self, docx_submission, student_b):
        # student_b requests student_a's submission -> 403
        r = requests.get(
            f"{BASE_URL}/api/submissions/{docx_submission['submission_id']}/preview-text",
            headers={"Authorization": f"Bearer {student_b['token']}"}, timeout=30,
        )
        assert r.status_code == 403, f"expected 403, got {r.status_code}: {r.text}"

    def test_student_can_preview_own_submission(self, docx_submission, student_a):
        r = requests.get(
            f"{BASE_URL}/api/submissions/{docx_submission['submission_id']}/preview-text",
            headers={"Authorization": f"Bearer {student_a['token']}"}, timeout=30,
        )
        assert r.status_code == 200, r.text
        assert "text" in r.json()

    def test_outsider_instructor_gets_403(self, docx_submission, outsider_instructor):
        r = requests.get(
            f"{BASE_URL}/api/submissions/{docx_submission['submission_id']}/preview-text",
            headers={"Authorization": f"Bearer {outsider_instructor['token']}"}, timeout=30,
        )
        assert r.status_code == 403, f"expected 403 for unrelated instructor, got {r.status_code}: {r.text}"


class TestPreviewTextLegacy410:
    """Legacy file_path-only submissions should return HTTP 410 with a clean message."""

    def test_legacy_submission_preview_text_returns_410(
        self, mongo, auth_headers, test_cohort, student_a, library_material
    ):
        fake_id = f"sub_legacy_prev_{uuid.uuid4().hex[:8]}"
        mongo.submissions.insert_one({
            "submission_id": fake_id,
            "material_id": library_material,
            "cohort_id": test_cohort,
            "student_id": student_a["user_id"],
            "file_path": "/app/backend/uploads/does_not_exist_preview.pdf",
            "gridfs_id": None,
            "file_name": "legacy.pdf",
            "status": "pending",
            "submitted_at": datetime.now(timezone.utc).isoformat(),
        })
        try:
            r = requests.get(
                f"{BASE_URL}/api/submissions/{fake_id}/preview-text",
                headers=auth_headers, timeout=30,
            )
            assert r.status_code == 410, f"expected 410, got {r.status_code}: {r.text}"
            # clean message body
            body = r.json()
            msg = (body.get("detail") or body.get("message") or "").lower()
            assert "no longer" in msg or "available" in msg or "missing" in msg or msg, \
                f"expected a clean 410 message, got: {body!r}"
        finally:
            mongo.submissions.delete_one({"submission_id": fake_id})


class TestSubmissionDetailRegression:
    """GET /api/submissions/{id} still returns material + student info."""

    def test_get_submission_returns_material_and_student(self, docx_submission, auth_headers):
        r = requests.get(
            f"{BASE_URL}/api/submissions/{docx_submission['submission_id']}",
            headers=auth_headers, timeout=30,
        )
        assert r.status_code == 200, r.text
        data = r.json()
        assert data["submission_id"] == docx_submission["submission_id"]
        assert data.get("material") is not None
        assert data.get("student") is not None
