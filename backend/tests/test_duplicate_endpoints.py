"""
Tests for library material + cohort duplication endpoints.

Endpoints under test:
- POST /api/library/materials/{material_id}/duplicate
- POST /api/cohorts/{cohort_id}/duplicate

Plus regression: /assign, /unassign, PUT /library/materials/{id}
"""
import io
import os
import uuid
from datetime import datetime, timezone, timedelta

import pytest
import requests
from pymongo import MongoClient
from bson import ObjectId
from docx import Document


BASE_URL = os.environ.get("REACT_APP_BACKEND_URL", "").rstrip("/")
MONGO_URL = os.environ.get("MONGO_URL", "mongodb://localhost:27017")
DB_NAME = os.environ.get("DB_NAME", "test_database")
SUPER_ADMIN_EMAIL = "slewis@theboostpad.org"


# ---------- fixtures ----------

@pytest.fixture(scope="module")
def mongo():
    client = MongoClient(MONGO_URL)
    yield client[DB_NAME]
    client.close()


def _seed_session(mongo, user_id: str, prefix: str = "test_dup") -> str:
    token = f"{prefix}_{uuid.uuid4().hex[:12]}"
    mongo.user_sessions.insert_one({
        "session_token": token,
        "user_id": user_id,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "expires_at": (datetime.now(timezone.utc) + timedelta(days=1)).isoformat(),
    })
    return token


@pytest.fixture(scope="module")
def admin_token(mongo):
    admin = mongo.users.find_one({"email": SUPER_ADMIN_EMAIL})
    assert admin, f"Super admin {SUPER_ADMIN_EMAIL} not seeded"
    token = _seed_session(mongo, admin["user_id"], "test_dup_admin")
    yield token
    mongo.user_sessions.delete_one({"session_token": token})


@pytest.fixture(scope="module")
def admin_headers(admin_token):
    return {"Authorization": f"Bearer {admin_token}"}


@pytest.fixture(scope="module")
def instructor_inside(mongo):
    """Instructor who will manage the source cohort."""
    uid = f"user_{uuid.uuid4().hex[:12]}"
    email = f"TEST_dup_inst_{uuid.uuid4().hex[:6]}@example.com"
    mongo.users.insert_one({
        "user_id": uid,
        "email": email,
        "name": "TEST Dup Instructor",
        "picture": None,
        "role": "instructor",
        "created_at": datetime.now(timezone.utc).isoformat(),
    })
    token = _seed_session(mongo, uid, "test_dup_inst")
    yield {"user_id": uid, "email": email, "token": token}
    mongo.user_sessions.delete_one({"session_token": token})
    mongo.users.delete_one({"user_id": uid})


@pytest.fixture(scope="module")
def instructor_outside(mongo):
    """Instructor NOT managing the source cohort (for ACL test)."""
    uid = f"user_{uuid.uuid4().hex[:12]}"
    email = f"TEST_dup_out_{uuid.uuid4().hex[:6]}@example.com"
    mongo.users.insert_one({
        "user_id": uid,
        "email": email,
        "name": "TEST Dup Outsider",
        "picture": None,
        "role": "instructor",
        "created_at": datetime.now(timezone.utc).isoformat(),
    })
    token = _seed_session(mongo, uid, "test_dup_out")
    yield {"user_id": uid, "email": email, "token": token}
    mongo.user_sessions.delete_one({"session_token": token})
    mongo.users.delete_one({"user_id": uid})


@pytest.fixture(scope="module")
def student(mongo):
    uid = f"user_{uuid.uuid4().hex[:12]}"
    email = f"TEST_dup_stu_{uuid.uuid4().hex[:6]}@example.com"
    mongo.users.insert_one({
        "user_id": uid,
        "email": email,
        "name": "TEST Dup Student",
        "picture": None,
        "role": "student",
        "created_at": datetime.now(timezone.utc).isoformat(),
    })
    yield {"user_id": uid, "email": email}
    mongo.users.delete_one({"user_id": uid})


def _make_docx_bytes(text: str) -> bytes:
    doc = Document()
    doc.add_heading("TEST Duplicate", level=1)
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture(scope="module")
def src_cohort(mongo, admin_headers, instructor_inside, student):
    """Create a cohort managed by instructor_inside, with student enrolled."""
    name = f"TEST_DUP_Cohort_{uuid.uuid4().hex[:6]}"
    r = requests.post(
        f"{BASE_URL}/api/cohorts",
        json={"name": name, "description": "duplicate test cohort"},
        headers=admin_headers, timeout=30,
    )
    assert r.status_code == 200, r.text
    cid = r.json()["cohort_id"]
    # Make instructor_inside manage it, add student, and release weeks 1+2
    mongo.cohorts.update_one(
        {"cohort_id": cid},
        {"$addToSet": {
            "instructor_ids": instructor_inside["user_id"],
            "student_ids": student["user_id"],
        },
         "$set": {"released_weeks": [1, 2]}},
    )
    yield {"cohort_id": cid, "name": name}
    # Teardown — best effort
    requests.delete(f"{BASE_URL}/api/cohorts/{cid}", headers=admin_headers, timeout=30)


@pytest.fixture(scope="module")
def library_material(mongo, admin_headers, src_cohort):
    """Library material (week=1, homework) assigned to src_cohort."""
    docx_bytes = _make_docx_bytes("Library homework reference content for duplication test")
    files = {"file": ("lib_hw.docx", docx_bytes,
                      "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    params = {
        "week_number": 1,
        "material_type": "homework",
        "title": f"TEST_DUP_Lib_HW_{uuid.uuid4().hex[:4]}",
        "description": "Library homework template",
    }
    r = requests.post(
        f"{BASE_URL}/api/library/materials",
        params=params, files=files, headers=admin_headers, timeout=60,
    )
    assert r.status_code == 200, r.text
    mid = r.json()["material_id"]

    # Assign to src cohort
    r2 = requests.post(
        f"{BASE_URL}/api/library/materials/{mid}/assign",
        json={"cohort_ids": [src_cohort["cohort_id"]]},
        headers=admin_headers, timeout=30,
    )
    assert r2.status_code == 200, r2.text

    yield {"material_id": mid, "bytes": docx_bytes, "title": params["title"]}

    requests.delete(f"{BASE_URL}/api/library/materials/{mid}", headers=admin_headers, timeout=30)


@pytest.fixture(scope="module")
def inline_material(mongo, admin_headers, src_cohort):
    """Cohort-specific (non-library) material attached to src_cohort, week 2."""
    docx_bytes = _make_docx_bytes("Cohort workbook content for duplication test")
    files = {"file": ("workbook.docx", docx_bytes,
                      "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    params = {
        "week_number": 2,
        "material_type": "workbook",
        "title": f"TEST_DUP_Inline_WB_{uuid.uuid4().hex[:4]}",
        "description": "Inline workbook",
    }
    r = requests.post(
        f"{BASE_URL}/api/cohorts/{src_cohort['cohort_id']}/materials",
        params=params, files=files, headers=admin_headers, timeout=60,
    )
    assert r.status_code == 200, r.text
    mid = r.json()["material_id"]
    yield {"material_id": mid, "bytes": docx_bytes, "title": params["title"]}


# ---------- tracker for created docs ----------
_created_cohort_ids: list = []
_created_material_ids: list = []


@pytest.fixture(scope="module", autouse=True)
def cleanup_created(mongo, admin_headers):
    yield
    # Best effort cleanup of any duplicated docs
    for cid in _created_cohort_ids:
        try:
            requests.delete(f"{BASE_URL}/api/cohorts/{cid}", headers=admin_headers, timeout=30)
        except Exception:
            pass
        mongo.cohorts.delete_one({"cohort_id": cid})
    for mid in _created_material_ids:
        # Drop any GridFS bytes if present
        m = mongo.materials.find_one({"material_id": mid})
        if m and m.get("gridfs_id"):
            try:
                mongo["fs.files"].delete_one({"_id": ObjectId(m["gridfs_id"])})
                mongo["fs.chunks"].delete_many({"files_id": ObjectId(m["gridfs_id"])})
            except Exception:
                pass
        mongo.materials.delete_one({"material_id": mid})


# ---------- Tests: Library material duplicate ----------

class TestLibraryDuplicate:

    def test_duplicate_library_material_basic(self, mongo, admin_headers, library_material, src_cohort):
        r = requests.post(
            f"{BASE_URL}/api/library/materials/{library_material['material_id']}/duplicate",
            headers=admin_headers, timeout=30,
        )
        assert r.status_code == 200, r.text
        body = r.json()
        new_id = body.get("material_id")
        assert new_id and new_id != library_material["material_id"]
        _created_material_ids.append(new_id)

        new_doc = mongo.materials.find_one({"material_id": new_id})
        src_doc = mongo.materials.find_one({"material_id": library_material["material_id"]})

        # Flag checks
        assert new_doc["is_library"] == True
        assert new_doc.get("cohort_ids", []) == []
        assert new_doc.get("duplicated_from") == library_material["material_id"]
        assert new_doc["title"].endswith(" (Copy)")
        assert new_doc["title"].startswith(library_material["title"])
        # New GridFS id, distinct
        assert new_doc.get("gridfs_id"), "duplicate must have new gridfs_id"
        assert new_doc["gridfs_id"] != src_doc["gridfs_id"], "gridfs_id must differ from source"

        # Source's cohort_ids should NOT be mutated by template duplication
        assert src_cohort["cohort_id"] in src_doc.get("cohort_ids", []), \
            "source library material should still be assigned to src_cohort"

    def test_duplicate_library_file_bytes_match(self, mongo, admin_headers, library_material):
        r = requests.post(
            f"{BASE_URL}/api/library/materials/{library_material['material_id']}/duplicate",
            headers=admin_headers, timeout=30,
        )
        assert r.status_code == 200, r.text
        new_id = r.json()["material_id"]
        _created_material_ids.append(new_id)

        # Download via API
        dl = requests.get(
            f"{BASE_URL}/api/materials/{new_id}/download",
            headers=admin_headers, timeout=30,
        )
        assert dl.status_code == 200, dl.text
        assert dl.content == library_material["bytes"], "duplicated file bytes must match source"

    def test_duplicate_legacy_library_returns_410(self, mongo, admin_headers):
        """Library material with only legacy file_path (no GridFS, missing on disk) → 410."""
        fake_mid = f"lib_legacy_{uuid.uuid4().hex[:8]}"
        mongo.materials.insert_one({
            "material_id": fake_mid,
            "is_library": True,
            "cohort_id": None,
            "cohort_ids": [],
            "week_number": 1,
            "material_type": "homework",
            "title": "TEST_DUP_legacy",
            "description": "legacy",
            "file_path": "/app/backend/uploads/does_not_exist.pdf",
            "gridfs_id": None,
            "file_name": "legacy.pdf",
            "uploaded_by": "system",
            "created_at": datetime.now(timezone.utc).isoformat(),
        })
        try:
            r = requests.post(
                f"{BASE_URL}/api/library/materials/{fake_mid}/duplicate",
                headers=admin_headers, timeout=30,
            )
            assert r.status_code == 410, f"expected 410, got {r.status_code}: {r.text}"
            assert "no longer available" in r.text.lower() or "cannot duplicate" in r.text.lower()
        finally:
            mongo.materials.delete_one({"material_id": fake_mid})


# ---------- Tests: Cohort duplicate ----------

class TestCohortDuplicate:

    def test_duplicate_cohort_basic(self, mongo, admin_headers, src_cohort,
                                    library_material, inline_material, instructor_inside):
        # Drop a fake submission on the source cohort to ensure it doesn't get copied
        fake_sub_id = f"sub_TEST_DUP_{uuid.uuid4().hex[:8]}"
        mongo.submissions.insert_one({
            "submission_id": fake_sub_id,
            "material_id": library_material["material_id"],
            "cohort_id": src_cohort["cohort_id"],
            "student_id": "ghost_student",
            "file_path": "",
            "gridfs_id": None,
            "file_name": "ghost.docx",
            "status": "pending",
            "submitted_at": datetime.now(timezone.utc).isoformat(),
        })

        try:
            r = requests.post(
                f"{BASE_URL}/api/cohorts/{src_cohort['cohort_id']}/duplicate",
                headers=admin_headers, timeout=60,
            )
            assert r.status_code == 200, r.text
            body = r.json()
            new_cid = body["cohort_id"]
            _created_cohort_ids.append(new_cid)

            new_cohort = mongo.cohorts.find_one({"cohort_id": new_cid})
            src_cohort_doc = mongo.cohorts.find_one({"cohort_id": src_cohort["cohort_id"]})

            # Name and template properties
            assert new_cohort["name"].endswith(" (Copy)")
            assert new_cohort["name"].startswith(src_cohort["name"])
            assert new_cohort.get("student_ids", []) == []
            admin_uid = mongo.users.find_one({"email": SUPER_ADMIN_EMAIL})["user_id"]
            assert admin_uid in new_cohort.get("instructor_ids", []), \
                "Current user must be in instructor_ids of duplicated cohort"
            # released_weeks carried over
            assert sorted(new_cohort.get("released_weeks", [])) == sorted(
                src_cohort_doc.get("released_weeks", [])
            )
            assert new_cohort.get("duplicated_from") == src_cohort["cohort_id"]

            # Library material now linked to BOTH cohorts
            lib_doc = mongo.materials.find_one({"material_id": library_material["material_id"]})
            assert src_cohort["cohort_id"] in lib_doc.get("cohort_ids", [])
            assert new_cid in lib_doc.get("cohort_ids", []), \
                "duplicated cohort must be added to library material's cohort_ids"
            assert body.get("library_materials_linked", 0) >= 1

            # Inline (cohort-specific) material is CLONED, not relinked
            new_inline_docs = list(mongo.materials.find({
                "cohort_id": new_cid,
                "is_library": {"$ne": True},
            }))
            assert len(new_inline_docs) >= 1, "inline materials should be cloned to new cohort"
            assert body.get("cohort_materials_copied", 0) >= 1

            inline_clone = next(
                (m for m in new_inline_docs if m.get("title") == inline_material["title"]),
                None,
            )
            assert inline_clone is not None, "Clone of inline material must exist"
            # Track for cleanup
            for m in new_inline_docs:
                _created_material_ids.append(m["material_id"])

            # Different material_id and different gridfs_id
            src_inline = mongo.materials.find_one({"material_id": inline_material["material_id"]})
            assert inline_clone["material_id"] != src_inline["material_id"]
            assert inline_clone.get("gridfs_id") and inline_clone["gridfs_id"] != src_inline["gridfs_id"]

            # File bytes match source
            dl = requests.get(
                f"{BASE_URL}/api/materials/{inline_clone['material_id']}/download",
                headers=admin_headers, timeout=30,
            )
            assert dl.status_code == 200, dl.text
            assert dl.content == inline_material["bytes"], \
                "cloned inline material file bytes must match source"

            # Submissions NOT copied
            new_subs = list(mongo.submissions.find({"cohort_id": new_cid}))
            assert new_subs == [], "submissions must not be copied to duplicated cohort"
        finally:
            mongo.submissions.delete_one({"submission_id": fake_sub_id})

    def test_duplicate_cohort_acl_outsider_403(self, mongo, src_cohort, instructor_outside):
        """An instructor not managing the source cohort should get 403."""
        headers = {"Authorization": f"Bearer {instructor_outside['token']}"}
        r = requests.post(
            f"{BASE_URL}/api/cohorts/{src_cohort['cohort_id']}/duplicate",
            headers=headers, timeout=30,
        )
        assert r.status_code == 403, f"expected 403, got {r.status_code}: {r.text}"

    def test_duplicate_cohort_404_for_missing(self, admin_headers):
        bogus = f"cohort_TEST_missing_{uuid.uuid4().hex[:6]}"
        r = requests.post(
            f"{BASE_URL}/api/cohorts/{bogus}/duplicate",
            headers=admin_headers, timeout=30,
        )
        assert r.status_code == 404, r.text


# ---------- Regression tests for existing endpoints ----------

class TestRegression:

    def test_assign_then_unassign(self, mongo, admin_headers, library_material, src_cohort):
        # Create a second cohort
        r = requests.post(
            f"{BASE_URL}/api/cohorts",
            json={"name": f"TEST_DUP_Cohort2_{uuid.uuid4().hex[:6]}", "description": "tmp"},
            headers=admin_headers, timeout=30,
        )
        assert r.status_code == 200, r.text
        c2 = r.json()["cohort_id"]
        try:
            r2 = requests.post(
                f"{BASE_URL}/api/library/materials/{library_material['material_id']}/assign",
                json={"cohort_ids": [c2]},
                headers=admin_headers, timeout=30,
            )
            assert r2.status_code == 200, r2.text
            doc = mongo.materials.find_one({"material_id": library_material["material_id"]})
            assert c2 in doc.get("cohort_ids", [])

            r3 = requests.post(
                f"{BASE_URL}/api/library/materials/{library_material['material_id']}/unassign",
                json={"cohort_id": c2},
                headers=admin_headers, timeout=30,
            )
            assert r3.status_code == 200, r3.text
            doc2 = mongo.materials.find_one({"material_id": library_material["material_id"]})
            assert c2 not in doc2.get("cohort_ids", [])
        finally:
            requests.delete(f"{BASE_URL}/api/cohorts/{c2}", headers=admin_headers, timeout=30)

    def test_update_library_material(self, mongo, admin_headers, library_material):
        new_title = f"TEST_DUP_Updated_{uuid.uuid4().hex[:4]}"
        r = requests.put(
            f"{BASE_URL}/api/library/materials/{library_material['material_id']}",
            params={"title": new_title, "description": "updated desc", "week_number": 1},
            headers=admin_headers, timeout=30,
        )
        assert r.status_code == 200, r.text
        doc = mongo.materials.find_one({"material_id": library_material["material_id"]})
        assert doc["title"] == new_title
        assert doc["description"] == "updated desc"
