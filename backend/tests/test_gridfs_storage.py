"""
End-to-end tests for GridFS file storage migration.
Verifies homework submissions + library materials are persisted to GridFS
and can be downloaded / reviewed without hitting disk.

Seeds a super-admin session and a test student directly in MongoDB to avoid
OAuth flows, then drives the public API via HTTP.
"""
import io
import os
import uuid
import time
from datetime import datetime, timezone, timedelta

import pytest
import requests
from pymongo import MongoClient
from bson import ObjectId
from docx import Document


BASE_URL = os.environ.get("REACT_APP_BACKEND_URL", "").rstrip("/")
MONGO_URL = os.environ.get("MONGO_URL", "mongodb://localhost:27017")
DB_NAME = os.environ.get("DB_NAME", "test_database")
SUPER_ADMIN_EMAIL = "slewis@theboostpad.org"


# ---------- fixtures ----------


@pytest.fixture(scope="module")
def mongo():
    client = MongoClient(MONGO_URL)
    yield client[DB_NAME]
    client.close()


@pytest.fixture(scope="module")
def admin_token(mongo):
    """Create a super-admin session token directly in Mongo."""
    admin = mongo.users.find_one({"email": SUPER_ADMIN_EMAIL})
    assert admin, f"Super admin {SUPER_ADMIN_EMAIL} not seeded in users collection"
    token = f"test_gridfs_{uuid.uuid4().hex[:12]}"
    mongo.user_sessions.insert_one({
        "session_token": token,
        "user_id": admin["user_id"],
        "created_at": datetime.now(timezone.utc).isoformat(),
        "expires_at": (datetime.now(timezone.utc) + timedelta(days=1)).isoformat(),
    })
    yield token
    mongo.user_sessions.delete_one({"session_token": token})


@pytest.fixture(scope="module")
def student(mongo):
    """Create a test student user + session directly in Mongo."""
    student_id = f"user_{uuid.uuid4().hex[:12]}"
    email = f"TEST_gridfs_{uuid.uuid4().hex[:6]}@example.com"
    mongo.users.insert_one({
        "user_id": student_id,
        "email": email,
        "name": "TEST GridFS Student",
        "picture": None,
        "role": "student",
        "created_at": datetime.now(timezone.utc).isoformat(),
    })
    token = f"test_gridfs_stu_{uuid.uuid4().hex[:12]}"
    mongo.user_sessions.insert_one({
        "session_token": token,
        "user_id": student_id,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "expires_at": (datetime.now(timezone.utc) + timedelta(days=1)).isoformat(),
    })
    yield {"user_id": student_id, "email": email, "token": token}
    mongo.user_sessions.delete_one({"session_token": token})
    mongo.users.delete_one({"user_id": student_id})


@pytest.fixture(scope="module")
def auth_headers(admin_token):
    return {"Authorization": f"Bearer {admin_token}"}


@pytest.fixture(scope="module")
def student_headers(student):
    return {"Authorization": f"Bearer {student['token']}"}


def _make_docx_bytes(text: str = "This is a TEST homework submission for GridFS verification.") -> bytes:
    doc = Document()
    doc.add_heading("TEST Homework", level=1)
    doc.add_paragraph(text)
    doc.add_paragraph("The student describes their leadership growth in detail " * 10)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture(scope="module")
def test_cohort(mongo, auth_headers):
    name = f"TEST_GridFS_Cohort_{uuid.uuid4().hex[:6]}"
    r = requests.post(
        f"{BASE_URL}/api/cohorts",
        json={"name": name, "description": "Automated GridFS regression"},
        headers=auth_headers,
        timeout=30,
    )
    assert r.status_code == 200, r.text
    cohort_id = r.json()["cohort_id"]
    yield cohort_id
    # Cleanup
    requests.delete(f"{BASE_URL}/api/cohorts/{cohort_id}", headers=auth_headers, timeout=30)


@pytest.fixture(scope="module")
def enrolled_student(mongo, test_cohort, student):
    mongo.cohorts.update_one(
        {"cohort_id": test_cohort},
        {"$addToSet": {"student_ids": student["user_id"]}},
    )
    return student


@pytest.fixture(scope="module")
def library_material(mongo, auth_headers, test_cohort):
    """Upload a homework material to the library and assign + release it."""
    docx_bytes = _make_docx_bytes("Homework assignment reference doc (TEST GridFS)")
    files = {"file": ("assignment.docx", docx_bytes,
                      "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    params = {
        "week_number": 1,
        "material_type": "homework",
        "title": f"TEST GridFS Homework {uuid.uuid4().hex[:4]}",
        "description": "Automated test material",
    }
    r = requests.post(
        f"{BASE_URL}/api/library/materials",
        params=params,
        files=files,
        headers=auth_headers,
        timeout=60,
    )
    assert r.status_code == 200, r.text
    material_id = r.json()["material_id"]

    # Verify doc has gridfs_id but no real file_path
    mat = mongo.materials.find_one({"material_id": material_id})
    assert mat is not None
    assert mat.get("gridfs_id"), "library material must have gridfs_id set"
    assert not mat.get("file_path"), "library material must NOT use disk file_path"

    # Assign + release
    r2 = requests.post(
        f"{BASE_URL}/api/library/materials/{material_id}/assign",
        json={"cohort_ids": [test_cohort]},
        headers=auth_headers, timeout=30,
    )
    assert r2.status_code == 200, r2.text
    r3 = requests.post(
        f"{BASE_URL}/api/cohorts/{test_cohort}/release-week",
        json={"week_number": 1},
        headers=auth_headers, timeout=30,
    )
    assert r3.status_code == 200, r3.text

    yield {"material_id": material_id, "bytes": docx_bytes}

    requests.delete(f"{BASE_URL}/api/library/materials/{material_id}",
                    headers=auth_headers, timeout=30)


# ---------- tests ----------


class TestLibraryMaterialGridFS:
    """Super admin library uploads persist to GridFS."""

    def test_material_upload_creates_gridfs_record(self, mongo, library_material):
        mat = mongo.materials.find_one({"material_id": library_material["material_id"]})
        assert mat["gridfs_id"]
        gridfs_file = mongo["fs.files"].find_one({"_id": ObjectId(mat["gridfs_id"])})
        assert gridfs_file is not None, "GridFS entry must exist for the uploaded material"
        assert gridfs_file["length"] == len(library_material["bytes"])

    def test_material_download_streams_gridfs_bytes(self, library_material, auth_headers):
        r = requests.get(
            f"{BASE_URL}/api/materials/{library_material['material_id']}/download",
            headers=auth_headers, timeout=30,
        )
        assert r.status_code == 200, r.text
        assert r.content == library_material["bytes"], "Downloaded bytes must match upload"


class TestSubmissionGridFS:
    """Student submission → GridFS → instructor download/review flow."""

    @pytest.fixture(scope="class")
    def submission(self, enrolled_student, library_material, student_headers, test_cohort, mongo):
        payload_bytes = _make_docx_bytes(
            "This submission contains the student's reflection on leadership. "
            "They discuss collaboration, accountability, and growth from Week 1."
        )
        files = {"file": ("my_homework.docx", payload_bytes,
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        r = requests.post(
            f"{BASE_URL}/api/materials/{library_material['material_id']}/submit",
            files=files,
            params={"cohort_id": test_cohort},
            headers=student_headers,
            timeout=60,
        )
        assert r.status_code == 200, r.text
        sub_doc = mongo.submissions.find_one({
            "material_id": library_material["material_id"],
            "student_id": enrolled_student["user_id"],
        })
        assert sub_doc is not None
        return {"submission_id": sub_doc["submission_id"], "bytes": payload_bytes, "doc": sub_doc}

    def test_submission_persisted_to_gridfs(self, mongo, submission):
        doc = mongo.submissions.find_one({"submission_id": submission["submission_id"]})
        assert doc["gridfs_id"], "submission must store gridfs_id"
        assert not doc.get("file_path"), "submission must NOT use disk file_path"
        gridfs_file = mongo["fs.files"].find_one({"_id": ObjectId(doc["gridfs_id"])})
        assert gridfs_file is not None
        assert gridfs_file["length"] == len(submission["bytes"])

    def test_instructor_download_submission_matches_bytes(self, submission, auth_headers):
        r = requests.get(
            f"{BASE_URL}/api/submissions/{submission['submission_id']}/download",
            headers=auth_headers, timeout=30,
        )
        assert r.status_code == 200, r.text
        assert r.content == submission["bytes"], "Streamed GridFS bytes must equal uploaded bytes"

    def test_get_submission_details_regression(self, submission, auth_headers):
        r = requests.get(
            f"{BASE_URL}/api/submissions/{submission['submission_id']}",
            headers=auth_headers, timeout=30,
        )
        assert r.status_code == 200, r.text
        data = r.json()
        assert data["submission_id"] == submission["submission_id"]
        assert data.get("material") is not None
        assert data.get("student") is not None

    def test_ai_review_reads_from_gridfs(self, submission, auth_headers):
        r = requests.post(
            f"{BASE_URL}/api/submissions/{submission['submission_id']}/review",
            headers=auth_headers, timeout=180,
        )
        assert r.status_code == 200, f"Review should succeed reading from GridFS: {r.status_code} {r.text}"
        body = r.json()
        # Should include AI-generated feedback text
        assert any(k in body for k in ("feedback", "ai_feedback", "message", "review")), body


class TestDeleteSubmissionRemovesGridFS:
    """Deleting a submission also removes its GridFS bytes."""

    def test_delete_submission_cleans_gridfs(
        self, mongo, enrolled_student, library_material, student_headers, test_cohort, auth_headers
    ):
        payload = _make_docx_bytes("TEST delete flow submission")
        files = {"file": ("delete_me.docx", payload,
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        # Fresh submission (replaces any previous — resubmission branch)
        r = requests.post(
            f"{BASE_URL}/api/materials/{library_material['material_id']}/submit",
            files=files, params={"cohort_id": test_cohort},
            headers=student_headers, timeout=60,
        )
        assert r.status_code == 200, r.text
        sub = mongo.submissions.find_one({
            "material_id": library_material["material_id"],
            "student_id": enrolled_student["user_id"],
        })
        gid = sub["gridfs_id"]
        assert mongo["fs.files"].find_one({"_id": ObjectId(gid)}) is not None

        r2 = requests.delete(
            f"{BASE_URL}/api/submissions/{sub['submission_id']}",
            headers=auth_headers, timeout=30,
        )
        assert r2.status_code == 200, r2.text
        # submission gone
        assert mongo.submissions.find_one({"submission_id": sub["submission_id"]}) is None
        # gridfs bytes gone
        assert mongo["fs.files"].find_one({"_id": ObjectId(gid)}) is None, \
            "GridFS file should be deleted along with submission"


class TestLegacyFilePathReturns410:
    """Legacy records stored only on disk should surface a clear 410 error."""

    def test_legacy_submission_returns_410(self, mongo, auth_headers, test_cohort, enrolled_student, library_material):
        fake_sub_id = f"sub_legacy_{uuid.uuid4().hex[:8]}"
        mongo.submissions.insert_one({
            "submission_id": fake_sub_id,
            "material_id": library_material["material_id"],
            "cohort_id": test_cohort,
            "student_id": enrolled_student["user_id"],
            "file_path": "/app/backend/uploads/does_not_exist.pdf",
            "gridfs_id": None,
            "file_name": "legacy.pdf",
            "status": "pending",
            "submitted_at": datetime.now(timezone.utc).isoformat(),
        })
        try:
            r = requests.get(
                f"{BASE_URL}/api/submissions/{fake_sub_id}/download",
                headers=auth_headers, timeout=30,
            )
            assert r.status_code == 410, f"expected 410 for legacy missing file, got {r.status_code}: {r.text}"
        finally:
            mongo.submissions.delete_one({"submission_id": fake_sub_id})
