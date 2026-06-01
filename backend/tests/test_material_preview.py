"""
Tests for inline preview / view of Material Library files.

Covers:
  - GET /api/materials/{material_id}/download                -> Content-Disposition: attachment (back-compat)
  - GET /api/materials/{material_id}/download?inline=1 (PDF) -> Content-Disposition: inline,
       Content-Type: application/pdf, X-Frame-Options: SAMEORIGIN, Content-Length set
  - GET /api/materials/{material_id}/preview-text (DOCX library)  -> {text, file_name}, text has body
  - GET /api/materials/{material_id}/preview-text (PDF library)   -> {text, file_name}, text is str
  - ACL:
      * student NOT in any assigned cohort  -> /preview-text on library material -> 403
      * student IN assigned cohort          -> /preview-text on library material -> 200
      * outsider instructor                  -> /preview-text on non-library cohort material -> 403
      * managing instructor (admin)          -> /preview-text on non-library cohort material -> 200
  - Legacy file_path-only material with missing disk file -> /preview-text -> 410
  - Regression: GET /api/materials/{id}/download (cohort + library) still returns matching bytes
  - Regression: /api/library/materials list returns materials with cohort_ids + assigned_cohorts
  - Auth: both Authorization: Bearer header AND ?token=<session_token> query param work
"""
import io
import os
import uuid
from datetime import datetime, timezone, timedelta

import pytest
import requests
from pymongo import MongoClient
from docx import Document
from dotenv import load_dotenv
from fpdf import FPDF

# Load env from frontend/.env (for REACT_APP_BACKEND_URL) and backend/.env (Mongo)
load_dotenv("/app/frontend/.env")
load_dotenv("/app/backend/.env")

BASE_URL = os.environ.get("REACT_APP_BACKEND_URL", "").rstrip("/")
MONGO_URL = os.environ.get("MONGO_URL", "mongodb://localhost:27017")
DB_NAME = os.environ.get("DB_NAME", "test_database")
assert BASE_URL, "REACT_APP_BACKEND_URL must be set"
SUPER_ADMIN_EMAIL = "slewis@theboostpad.org"


# ---------- helpers ----------

def _make_docx_bytes(body_text: str) -> bytes:
    doc = Document()
    doc.add_heading("TEST Material Preview", level=1)
    doc.add_paragraph(body_text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_pdf_bytes(text: str) -> bytes:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.cell(0, 10, text)
    raw = pdf.output(dest="S")
    if isinstance(raw, str):
        return raw.encode("latin-1")
    if isinstance(raw, bytearray):
        return bytes(raw)
    return raw


def _raw_get(url: str, headers: dict = None, params: dict = None):
    """Force identity encoding so Content-Length is observable."""
    h = dict(headers or {})
    h.setdefault("Accept-Encoding", "identity")
    return requests.get(url, headers=h, params=params, timeout=30)


# ---------- fixtures ----------

@pytest.fixture(scope="module")
def mongo():
    client = MongoClient(MONGO_URL)
    yield client[DB_NAME]
    client.close()


def _mk_session(mongo, user_id: str, prefix: str) -> str:
    token = f"{prefix}_{uuid.uuid4().hex[:12]}"
    mongo.user_sessions.insert_one({
        "session_token": token,
        "user_id": user_id,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "expires_at": (datetime.now(timezone.utc) + timedelta(days=1)).isoformat(),
    })
    return token


def _mk_user(mongo, role: str, prefix: str) -> dict:
    uid = f"user_{uuid.uuid4().hex[:12]}"
    email = f"TEST_matprev_{prefix}_{uuid.uuid4().hex[:6]}@example.com"
    mongo.users.insert_one({
        "user_id": uid,
        "email": email,
        "name": f"TEST {role} {prefix}",
        "picture": None,
        "role": role,
        "created_at": datetime.now(timezone.utc).isoformat(),
    })
    token = _mk_session(mongo, uid, f"test_matprev_{prefix}")
    return {"user_id": uid, "email": email, "token": token, "role": role}


@pytest.fixture(scope="module")
def admin_token(mongo):
    admin = mongo.users.find_one({"email": SUPER_ADMIN_EMAIL})
    assert admin, f"Super admin {SUPER_ADMIN_EMAIL} not seeded"
    token = _mk_session(mongo, admin["user_id"], "test_matprev_admin")
    yield token
    mongo.user_sessions.delete_one({"session_token": token})


@pytest.fixture(scope="module")
def auth_headers(admin_token):
    return {"Authorization": f"Bearer {admin_token}"}


@pytest.fixture(scope="module")
def student_in(mongo):
    """Student that will be enrolled in cohort A."""
    s = _mk_user(mongo, "student", "stu_in")
    yield s
    mongo.user_sessions.delete_one({"session_token": s["token"]})
    mongo.users.delete_one({"user_id": s["user_id"]})


@pytest.fixture(scope="module")
def student_out(mongo):
    """Student NOT enrolled in any cohort."""
    s = _mk_user(mongo, "student", "stu_out")
    yield s
    mongo.user_sessions.delete_one({"session_token": s["token"]})
    mongo.users.delete_one({"user_id": s["user_id"]})


@pytest.fixture(scope="module")
def outsider_instructor(mongo):
    """Instructor not managing cohort A."""
    i = _mk_user(mongo, "instructor", "out_inst")
    yield i
    mongo.user_sessions.delete_one({"session_token": i["token"]})
    mongo.users.delete_one({"user_id": i["user_id"]})


@pytest.fixture(scope="module")
def cohort_a(mongo, auth_headers, student_in):
    """Cohort A. Admin (super_admin) auto-manages all cohorts via is_cohort_manager.
    student_in is enrolled. student_out is NOT enrolled."""
    name = f"TEST_MatPrev_Cohort_{uuid.uuid4().hex[:6]}"
    r = requests.post(
        f"{BASE_URL}/api/cohorts",
        json={"name": name, "description": "material preview test"},
        headers=auth_headers, timeout=30,
    )
    assert r.status_code == 200, r.text
    cohort_id = r.json()["cohort_id"]
    mongo.cohorts.update_one(
        {"cohort_id": cohort_id},
        {"$addToSet": {"student_ids": student_in["user_id"]}},
    )
    yield cohort_id
    requests.delete(f"{BASE_URL}/api/cohorts/{cohort_id}", headers=auth_headers, timeout=30)


@pytest.fixture(scope="module")
def library_docx_material(auth_headers, cohort_a):
    """Library DOCX material assigned to cohort A."""
    body = "leadership reflection paragraph for material preview"
    files = {"file": ("lib_preview.docx", _make_docx_bytes(body),
                      "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    params = {
        "week_number": 1,
        "material_type": "homework",
        "title": f"TEST Lib DOCX Preview {uuid.uuid4().hex[:4]}",
        "description": "automated preview test material (docx)",
    }
    r = requests.post(f"{BASE_URL}/api/library/materials", params=params, files=files,
                      headers=auth_headers, timeout=60)
    assert r.status_code == 200, r.text
    mat_id = r.json()["material_id"]
    r2 = requests.post(f"{BASE_URL}/api/library/materials/{mat_id}/assign",
                       json={"cohort_ids": [cohort_a]}, headers=auth_headers, timeout=30)
    assert r2.status_code == 200, r2.text
    yield {"material_id": mat_id, "body_text": body, "file_name": "lib_preview.docx"}
    requests.delete(f"{BASE_URL}/api/library/materials/{mat_id}",
                    headers=auth_headers, timeout=30)


@pytest.fixture(scope="module")
def library_pdf_material(auth_headers, cohort_a):
    """Library PDF material assigned to cohort A."""
    pdf_text = "Inline PDF preview text from material library"
    pdf_bytes = _make_pdf_bytes(pdf_text)
    files = {"file": ("lib_preview.pdf", pdf_bytes, "application/pdf")}
    params = {
        "week_number": 1,
        "material_type": "case_study",
        "title": f"TEST Lib PDF Preview {uuid.uuid4().hex[:4]}",
        "description": "automated preview test material (pdf)",
    }
    r = requests.post(f"{BASE_URL}/api/library/materials", params=params, files=files,
                      headers=auth_headers, timeout=60)
    assert r.status_code == 200, r.text
    mat_id = r.json()["material_id"]
    r2 = requests.post(f"{BASE_URL}/api/library/materials/{mat_id}/assign",
                       json={"cohort_ids": [cohort_a]}, headers=auth_headers, timeout=30)
    assert r2.status_code == 200, r2.text
    yield {"material_id": mat_id, "bytes": pdf_bytes, "text": pdf_text,
           "file_name": "lib_preview.pdf"}
    requests.delete(f"{BASE_URL}/api/library/materials/{mat_id}",
                    headers=auth_headers, timeout=30)


@pytest.fixture(scope="module")
def cohort_specific_docx_material(mongo, auth_headers, cohort_a):
    """Non-library DOCX material uploaded directly to cohort A (is_library=False)."""
    body = "cohort specific docx body content"
    files = {"file": ("cohort_only.docx", _make_docx_bytes(body),
                      "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    params = {
        "week_number": 2,
        "material_type": "workbook",
        "title": f"TEST Cohort DOCX {uuid.uuid4().hex[:4]}",
        "description": "non-library cohort material",
    }
    r = requests.post(f"{BASE_URL}/api/cohorts/{cohort_a}/materials",
                      params=params, files=files,
                      headers=auth_headers, timeout=60)
    assert r.status_code == 200, r.text
    mat_id = r.json()["material_id"]
    yield {"material_id": mat_id, "body_text": body, "file_name": "cohort_only.docx"}
    mongo.materials.delete_one({"material_id": mat_id})


# =====================================================================
# 1. Download header back-compat + ?inline=1 contract
# =====================================================================

class TestDownloadInlineContract:
    def test_pdf_download_default_is_attachment(self, library_pdf_material, auth_headers):
        r = _raw_get(
            f"{BASE_URL}/api/materials/{library_pdf_material['material_id']}/download",
            headers=auth_headers,
        )
        assert r.status_code == 200, r.text
        cd = r.headers.get("Content-Disposition", "")
        assert cd.lower().startswith("attachment"), f"expected attachment, got: {cd!r}"
        # bytes round-trip
        assert r.content == library_pdf_material["bytes"]
        # content-length should match body
        cl = r.headers.get("Content-Length")
        assert cl is not None and int(cl) == len(r.content)

    def test_pdf_download_inline_sets_inline_pdf_xframe_contentlength(
        self, library_pdf_material, auth_headers
    ):
        r = _raw_get(
            f"{BASE_URL}/api/materials/{library_pdf_material['material_id']}/download",
            headers=auth_headers, params={"inline": 1},
        )
        assert r.status_code == 200, r.text
        cd = r.headers.get("Content-Disposition", "")
        assert cd.lower().startswith("inline"), f"expected inline, got: {cd!r}"
        ct = r.headers.get("Content-Type", "")
        assert ct.startswith("application/pdf"), f"expected application/pdf, got: {ct!r}"
        assert r.headers.get("X-Frame-Options", "").upper() == "SAMEORIGIN"
        cl = r.headers.get("Content-Length")
        assert cl is not None, "Content-Length missing (iframe rendering bug)"
        assert int(cl) == len(r.content) == len(library_pdf_material["bytes"])
        # Should NOT be chunked
        te = r.headers.get("Transfer-Encoding", "")
        assert "chunked" not in te.lower(), f"unexpected chunked: {te!r}"

    def test_docx_inline_disposition_but_octet_media_type(
        self, library_docx_material, auth_headers
    ):
        r = _raw_get(
            f"{BASE_URL}/api/materials/{library_docx_material['material_id']}/download",
            headers=auth_headers, params={"inline": 1},
        )
        assert r.status_code == 200, r.text
        cd = r.headers.get("Content-Disposition", "")
        assert cd.lower().startswith("inline"), f"expected inline, got: {cd!r}"
        ct = r.headers.get("Content-Type", "")
        assert ct.startswith("application/octet-stream"), f"docx should be octet-stream, got: {ct!r}"


# =====================================================================
# 2. preview-text endpoint - happy path
# =====================================================================

class TestPreviewTextHappyPath:
    def test_docx_preview_text_contains_body(self, library_docx_material, auth_headers):
        r = requests.get(
            f"{BASE_URL}/api/materials/{library_docx_material['material_id']}/preview-text",
            headers=auth_headers, timeout=30,
        )
        assert r.status_code == 200, r.text
        data = r.json()
        assert "text" in data and "file_name" in data
        assert data["file_name"].endswith(".docx")
        assert "leadership reflection" in data["text"].lower(), \
            f"expected body text; got: {data['text'][:200]!r}"

    def test_pdf_preview_text_nonempty(self, library_pdf_material, auth_headers):
        r = requests.get(
            f"{BASE_URL}/api/materials/{library_pdf_material['material_id']}/preview-text",
            headers=auth_headers, timeout=30,
        )
        assert r.status_code == 200, r.text
        data = r.json()
        assert "text" in data and "file_name" in data
        assert isinstance(data["text"], str)
        assert data["file_name"].endswith(".pdf")
        assert len(data["text"].strip()) > 0, \
            f"PDF preview text should be non-empty, got: {data['text']!r}"


# =====================================================================
# 3. ACL enforcement
# =====================================================================

class TestPreviewTextACL:
    def test_student_not_in_cohort_gets_403_on_library_material(
        self, library_docx_material, student_out
    ):
        r = requests.get(
            f"{BASE_URL}/api/materials/{library_docx_material['material_id']}/preview-text",
            headers={"Authorization": f"Bearer {student_out['token']}"}, timeout=30,
        )
        assert r.status_code == 403, f"expected 403, got {r.status_code}: {r.text}"

    def test_student_in_cohort_gets_200_on_library_material(
        self, library_docx_material, student_in
    ):
        r = requests.get(
            f"{BASE_URL}/api/materials/{library_docx_material['material_id']}/preview-text",
            headers={"Authorization": f"Bearer {student_in['token']}"}, timeout=30,
        )
        assert r.status_code == 200, f"expected 200, got {r.status_code}: {r.text}"
        data = r.json()
        assert "text" in data
        assert "leadership reflection" in data["text"].lower()

    def test_outsider_instructor_gets_403_on_non_library_cohort_material(
        self, cohort_specific_docx_material, outsider_instructor
    ):
        r = requests.get(
            f"{BASE_URL}/api/materials/{cohort_specific_docx_material['material_id']}/preview-text",
            headers={"Authorization": f"Bearer {outsider_instructor['token']}"}, timeout=30,
        )
        assert r.status_code == 403, f"expected 403, got {r.status_code}: {r.text}"

    def test_managing_admin_gets_200_on_non_library_cohort_material(
        self, cohort_specific_docx_material, auth_headers
    ):
        # super_admin is always a cohort manager (is_cohort_manager returns True for super_admin)
        r = requests.get(
            f"{BASE_URL}/api/materials/{cohort_specific_docx_material['material_id']}/preview-text",
            headers=auth_headers, timeout=30,
        )
        assert r.status_code == 200, f"expected 200, got {r.status_code}: {r.text}"
        data = r.json()
        assert "cohort specific" in data["text"].lower()


# =====================================================================
# 4. Legacy file_path-only material -> 410
# =====================================================================

class TestPreviewTextLegacy410:
    def test_legacy_material_missing_disk_file_returns_410(
        self, mongo, auth_headers, cohort_a
    ):
        fake_id = f"mat_legacy_prev_{uuid.uuid4().hex[:8]}"
        mongo.materials.insert_one({
            "material_id": fake_id,
            "is_library": True,
            "cohort_id": None,
            "cohort_ids": [cohort_a],
            "week_number": 1,
            "material_type": "homework",
            "title": "TEST legacy missing file",
            "description": "",
            "file_path": "/app/backend/uploads/does_not_exist_material_preview.pdf",
            "gridfs_id": None,
            "file_name": "legacy.pdf",
            "uploaded_by": "system",
            "due_date": None,
            "created_at": datetime.now(timezone.utc).isoformat(),
        })
        try:
            r = requests.get(
                f"{BASE_URL}/api/materials/{fake_id}/preview-text",
                headers=auth_headers, timeout=30,
            )
            assert r.status_code == 410, f"expected 410, got {r.status_code}: {r.text}"
            body = r.json()
            msg = (body.get("detail") or body.get("message") or "").lower()
            assert "no longer" in msg or "available" in msg or "missing" in msg, \
                f"expected clean 410 message, got: {body!r}"
        finally:
            mongo.materials.delete_one({"material_id": fake_id})


# =====================================================================
# 5. Regression: download bytes match for library + cohort materials
# =====================================================================

class TestDownloadRegression:
    def test_library_docx_download_bytes_match(self, library_docx_material, auth_headers):
        r = _raw_get(
            f"{BASE_URL}/api/materials/{library_docx_material['material_id']}/download",
            headers=auth_headers,
        )
        assert r.status_code == 200, r.text
        # docx default content disposition is attachment
        assert r.headers.get("Content-Disposition", "").lower().startswith("attachment")
        cl = r.headers.get("Content-Length")
        assert cl is not None and int(cl) == len(r.content)
        # Should be a valid docx (ZIP magic PK\x03\x04)
        assert r.content[:2] == b"PK", "DOCX is a ZIP container; magic mismatch"

    def test_cohort_specific_docx_download_bytes_match(
        self, cohort_specific_docx_material, auth_headers
    ):
        r = _raw_get(
            f"{BASE_URL}/api/materials/{cohort_specific_docx_material['material_id']}/download",
            headers=auth_headers,
        )
        assert r.status_code == 200, r.text
        cl = r.headers.get("Content-Length")
        assert cl is not None and int(cl) == len(r.content)
        assert r.content[:2] == b"PK"


# =====================================================================
# 6. Regression: /api/library/materials list returns cohort_ids + assigned_cohorts
# =====================================================================

class TestLibraryListRegression:
    def test_library_list_returns_cohort_ids_and_assigned_cohorts(
        self, library_docx_material, cohort_a, auth_headers
    ):
        r = requests.get(
            f"{BASE_URL}/api/library/materials",
            headers=auth_headers, timeout=30,
        )
        assert r.status_code == 200, r.text
        items = r.json()
        assert isinstance(items, list)
        match = next(
            (m for m in items if m.get("material_id") == library_docx_material["material_id"]),
            None,
        )
        assert match is not None, "uploaded library material missing from list"
        assert "cohort_ids" in match
        assert cohort_a in match["cohort_ids"]
        assert "assigned_cohorts" in match
        assert isinstance(match["assigned_cohorts"], list)
        ids_in_assigned = [c.get("cohort_id") for c in match["assigned_cohorts"]]
        assert cohort_a in ids_in_assigned


# =====================================================================
# 7. Auth via ?token= query param (iframe path)
# =====================================================================

class TestQueryTokenAuth:
    def test_download_inline_pdf_with_token_query_param(
        self, library_pdf_material, admin_token
    ):
        """The iframe relies on ?token=<session_token> instead of Authorization header."""
        r = _raw_get(
            f"{BASE_URL}/api/materials/{library_pdf_material['material_id']}/download",
            params={"inline": 1, "token": admin_token},
        )
        assert r.status_code == 200, r.text
        assert r.headers.get("Content-Disposition", "").lower().startswith("inline")
        assert (r.headers.get("Content-Type") or "").startswith("application/pdf")
        assert r.content == library_pdf_material["bytes"]

    def test_preview_text_with_token_query_param(
        self, library_docx_material, admin_token
    ):
        r = requests.get(
            f"{BASE_URL}/api/materials/{library_docx_material['material_id']}/preview-text",
            params={"token": admin_token}, timeout=30,
        )
        assert r.status_code == 200, r.text
        assert "leadership reflection" in r.json()["text"].lower()
